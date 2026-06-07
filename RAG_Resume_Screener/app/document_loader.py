import os
from pathlib import Path
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document


class DocumentLoader:
    @staticmethod
    def load_document(file_path: str) -> list[Document]:
        """Load document with proper page tracking across PDF, DOCX, and TXT files."""
        path = Path(file_path)
        file_extension = path.suffix.lower()

        if not path.exists():
            raise FileNotFoundError(f"The file at {file_path} does not exist.")

        if file_extension == ".pdf":
            loader = PyPDFLoader(str(path))
            documents = loader.load()
            for i, doc in enumerate(documents):
                doc.metadata["page"] = i + 1
                doc.metadata["source_type"] = "pdf"
            return documents

        elif file_extension == ".docx":
            from docx import Document as DocxReader
            doc_obj = DocxReader(str(path))
            documents = []
            words_per_page = 500
            current_page = 1
            current_content = []
            word_count = 0

            for para in doc_obj.paragraphs:
                if para.text.strip():
                    para_words = len(para.text.split())
                    if word_count + para_words > words_per_page:
                        if current_content:
                            documents.append(Document(
                                page_content="\n".join(current_content),
                                metadata={
                                    "source": str(path),
                                    "page": current_page,
                                    "source_type": "docx"
                                }
                            ))
                            current_content = []
                            word_count = 0
                            current_page += 1
                    current_content.append(para.text)
                    word_count += para_words

            if current_content:
                documents.append(Document(
                    page_content="\n".join(current_content),
                    metadata={
                        "source": str(path),
                        "page": current_page,
                        "source_type": "docx"
                    }
                ))
            return documents

        elif file_extension == ".txt":
            loader = TextLoader(str(path), encoding="utf-8")
            documents = loader.load()
            for doc in documents:
                doc.metadata["page"] = 1
                doc.metadata["source_type"] = "txt"
            return documents

        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
